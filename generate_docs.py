#!/usr/bin/env python3
"""
Generate all 8 GroomGuru documentation .docx files.

Usage:
  python generate_docs.py [--be-path PATH] [--fe-path PATH] [--output PATH]

Defaults:
  --be-path ../groomguru-service
  --fe-path ../groomguru-ui
  --output  ./docs
"""

import argparse
import os
import glob

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_args():
    p = argparse.ArgumentParser(description="Generate GroomGuru documentation")
    p.add_argument("--be-path", default="../groomguru-service", help="Path to groomguru-service repo")
    p.add_argument("--fe-path", default="../groomguru-ui", help="Path to groomguru-ui repo")
    p.add_argument("--output", default="./docs", help="Output directory for .docx files")
    return p.parse_args()


# ── Helpers ──────────────────────────────────────────

def styled_doc(title):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    heading_style = doc.styles['Heading 1']
    heading_style.font.color.rgb = RGBColor(55, 48, 163)
    heading_style.font.size = Pt(22)
    heading_style.font.bold = True

    h2 = doc.styles['Heading 2']
    h2.font.color.rgb = RGBColor(79, 70, 229)
    h2.font.size = Pt(16)

    h3 = doc.styles['Heading 3']
    h3.font.color.rgb = RGBColor(99, 102, 241)
    h3.font.size = Pt(13)

    for _ in range(4):
        doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run("GroomGuru")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(55, 48, 163)

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = tp2.add_run(title)
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    tp3 = doc.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = tp3.add_run("AI-Powered Interview Preparation Platform")
    run3.font.size = Pt(12)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()
    return doc


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_code(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("Note: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 120, 0)
    run2 = p.add_run(text)
    run2.font.italic = True


def save(doc, output_dir, name):
    path = os.path.join(output_dir, name)
    doc.save(path)
    print(f"  Created: {path}")


def read_file(path):
    """Read a file if it exists, return its content or a placeholder."""
    if os.path.isfile(path):
        with open(path, 'r', errors='replace') as f:
            return f.read()
    return f"[File not found: {path}]"


def count_files(directory, ext):
    """Count files with given extension recursively."""
    return len(glob.glob(os.path.join(directory, '**', f'*{ext}'), recursive=True))


# ── Document generators ──────────────────────────────

def doc1(output_dir, be_path, fe_path):
    doc = styled_doc("Project Overview & Features")

    doc.add_heading("1. What is GroomGuru?", level=1)
    doc.add_paragraph(
        "GroomGuru is a web application that helps people prepare for job interviews using Artificial Intelligence (AI). "
        "Think of it as your personal interview coach that is available 24/7. Instead of asking a friend to practice "
        "interview questions with you, GroomGuru uses AI to act as the interviewer — it asks you questions, listens to "
        "your answers, and at the end gives you a detailed report on how you did."
    )
    doc.add_paragraph(
        "The name 'GroomGuru' comes from 'Grooming' (preparing/polishing yourself) + 'Guru' (teacher/expert). "
        "It literally means: an expert that helps you groom yourself for interviews."
    )

    doc.add_heading("2. What Problem Does It Solve?", level=1)
    doc.add_paragraph("Interview preparation is hard because:")
    for b in [
        "You don't always have someone available to practice with",
        "You don't know what questions to expect for your specific role",
        "You can't get honest, unbiased feedback on your answers",
        "You can't practice speaking confidently under pressure",
        "You don't know if you're using too many filler words (um, uh, like)",
    ]:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph(
        "GroomGuru solves all of these problems by providing an AI interviewer that adapts to your role, "
        "experience level, and industry."
    )

    doc.add_heading("3. Key Features", level=1)

    doc.add_heading("3.1 Two Interview Modes", level=2)
    doc.add_paragraph(
        "Voice Interview: You speak into your microphone. The AI listens to your voice in real-time using "
        "OpenAI's Realtime API and responds with its own AI-generated voice. It feels like a real phone or "
        "video interview. The conversation flows naturally — you speak, the AI responds, you answer back."
    )
    doc.add_paragraph(
        "Chat Interview: You type your answers. The AI responds with text. This is useful if you're in a "
        "noisy environment or prefer typing. It's like chatting with an interviewer over messaging."
    )

    doc.add_heading("3.2 Smart Interview Configuration", level=2)
    doc.add_paragraph("Before starting an interview, you configure:")
    for item, desc in [
        ("Role", "What job are you interviewing for? (e.g., Software Engineer, Product Manager, Data Scientist)"),
        ("Experience Level", "Are you a Fresher, Junior, Mid-level, Senior, Lead, or Principal?"),
        ("Interview Type", "Technical, Behavioral, System Design, HR/Culture Fit, or Mixed"),
        ("Duration", "How long? 5, 10, 15, 20, or 30 minutes"),
        ("Company", "Optionally specify the company to tailor questions"),
        ("Simulation Mode", "Human mode (natural conversation) or Bot mode (rapid-fire detection style)"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{item}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("3.3 CV/Resume Upload & Auto-fill", level=2)
    doc.add_paragraph(
        "You can upload your CV (resume) as a PDF file. The AI reads your CV and automatically fills in your "
        "profile details like current role, skills, experience level, industry, and even your phone number. "
        "This saves you time and also gives the AI interviewer more context about you, making the interview "
        "questions more relevant and personalized."
    )

    doc.add_heading("3.4 Real-time Speech Analysis", level=2)
    doc.add_paragraph(
        "During voice interviews, the app tracks your speech patterns. After the interview, you get a Speech "
        "Analysis report that includes: words per minute, filler word count (um, uh, like, you know), average "
        "response length, vocabulary diversity score, and confidence indicators."
    )

    doc.add_heading("3.5 Proctor Mode (Optional)", level=2)
    doc.add_paragraph(
        "In voice interviews, you can enable 'Proctor Mode'. This uses your webcam and MediaPipe AI to track "
        "where you're looking. If you look away from the screen too much (suggesting you might be reading "
        "answers from somewhere), it flags this in the report. This simulates a proctored interview experience."
    )

    doc.add_heading("3.6 Comprehensive Hiring Report", level=2)
    doc.add_paragraph(
        "After every interview, the AI generates a detailed report that evaluates your performance. This report "
        "includes scores for technical knowledge, communication, problem-solving, and overall impression. "
        "It also provides specific feedback on what you did well and what you can improve."
    )

    doc.add_heading("3.7 User Profile & Account Management", level=2)
    doc.add_paragraph(
        "A complete profile system with: first/last name, email, phone number (with country code picker that "
        "fetches all world country codes from an API), profile avatar (upload or choose gradient), password "
        "management, and the ability to soft-delete your account."
    )

    doc.add_heading("3.8 Authentication", level=2)
    doc.add_paragraph(
        "Two ways to sign in: (1) Email + Password — traditional registration with password strength validation "
        "(uppercase, lowercase, number, special character, 8+ characters), or (2) Google Sign-In — one-click "
        "login using your Google account via OAuth 2.0."
    )

    doc.add_heading("3.9 Idle Timer & Auto Logout", level=2)
    doc.add_paragraph(
        "For security, if you leave the app idle for 15 minutes, it shows a warning. If you don't respond "
        "within 60 seconds, it automatically logs you out. This is paused during active interviews."
    )

    doc.add_heading("4. User Flow (Step by Step)", level=1)
    for i, step in enumerate([
        "User opens the app and sees the login page",
        "Signs up with email/password OR logs in with Google",
        "Lands on Home Dashboard",
        "Goes to Profile, uploads CV, AI auto-fills profile details",
        "Goes to Practice Interview, chooses Voice or Chat",
        "Configures interview settings (role, type, duration, etc.)",
        "Starts interview — AI asks questions and user responds",
        "Timer runs out or user ends interview manually",
        "AI generates comprehensive Hiring Report",
        "User reviews report with speech analysis, scores, and feedback",
        "User can go back and practice again",
    ], 1):
        doc.add_paragraph(f"Step {i}: {step}", style='List Number')

    doc.add_heading("5. Codebase Statistics", level=1)
    be_java = count_files(os.path.join(be_path, 'src'), '.java')
    fe_tsx = count_files(os.path.join(fe_path, 'src'), '.tsx')
    fe_ts = count_files(os.path.join(fe_path, 'src'), '.ts')
    add_table(doc, ["Metric", "Value"], [
        ["Backend Java files", str(be_java)],
        ["Frontend TSX components", str(fe_tsx)],
        ["Frontend TS files", str(fe_ts)],
        ["Database tables", "3 (users, user_profiles, interview_profiles)"],
        ["REST API endpoints", "19"],
        ["WebSocket endpoints", "1 (/ws/realtime)"],
    ])

    save(doc, output_dir, "01_Project_Overview_and_Features.docx")


def doc2(output_dir, be_path, fe_path):
    doc = styled_doc("Tech Stack & Architecture")

    doc.add_heading("1. What is a 'Tech Stack'?", level=1)
    doc.add_paragraph(
        "A tech stack is the collection of technologies, programming languages, frameworks, and tools used "
        "to build an application. Think of it like the ingredients list for a recipe — each ingredient "
        "(technology) plays a specific role in making the final dish (application)."
    )
    doc.add_paragraph(
        "GroomGuru has two main parts: the Frontend (what users see and interact with in their browser) "
        "and the Backend (the server that processes data, talks to AI, and manages the database). These "
        "two parts communicate over the internet using APIs (Application Programming Interfaces)."
    )

    doc.add_heading("2. Frontend Technologies", level=1)
    for name, desc in [
        ("React 19", "A JavaScript library for building user interfaces. It lets you create reusable UI components (buttons, forms, cards) and updates the page efficiently when data changes."),
        ("TypeScript 5.9", "A superset of JavaScript that adds 'types' — it helps catch errors before the code runs."),
        ("Vite 7.3", "The build tool and development server. It takes all your code files and bundles them into optimized files the browser can understand."),
        ("Tailwind CSS v4", "A CSS framework that lets you style elements directly in the HTML using utility classes."),
        ("React Router v7", "Handles navigation between pages without reloading (Single Page Application)."),
        ("Axios", "A library for making HTTP requests to the backend server."),
        ("@react-oauth/google", "Handles Google Sign-In OAuth flow."),
        ("@mediapipe/tasks-vision", "Google's AI library for face landmark detection in Proctor Mode."),
    ]:
        doc.add_heading(f"  {name}", level=3)
        doc.add_paragraph(desc)

    doc.add_heading("3. Backend Technologies", level=1)
    for name, desc in [
        ("Java 21", "The programming language for the backend. Widely used in enterprise applications."),
        ("Spring Boot 4.0.3", "Framework that handles web requests, security, database connections, and more."),
        ("Spring Security", "Handles authentication (JWT tokens) and authorization."),
        ("Spring Data JPA + Hibernate", "Manages database interactions using Java objects instead of raw SQL."),
        ("PostgreSQL 14", "The relational database storing all user and profile data."),
        ("Flyway", "Database migration tool for versioned, repeatable schema changes."),
        ("OpenAI API (GPT-4o, GPT-4o-mini)", "AI engine for interview questions, CV parsing, and report generation."),
        ("OpenAI Whisper", "Speech-to-text transcription for voice interviews."),
        ("OpenAI TTS", "Text-to-speech for AI voice responses."),
        ("Apache PDFBox 3.0", "Extracts text from uploaded PDF CVs."),
        ("JJWT 0.12.6", "Creates and validates JWT authentication tokens."),
        ("Google API Client", "Verifies Google OAuth tokens on the backend."),
        ("Spring WebSocket", "Real-time two-way communication for voice interviews."),
    ]:
        doc.add_heading(f"  {name}", level=3)
        doc.add_paragraph(desc)

    doc.add_heading("4. How Frontend and Backend Communicate", level=1)
    doc.add_heading("4.1 REST API (Request-Response)", level=2)
    doc.add_paragraph("For most operations, the frontend sends an HTTP request and receives a JSON response.")
    add_code(doc, "Frontend (port 5173) ---HTTP Request---> Backend (port 8080)\nFrontend <---JSON Response--- Backend")

    doc.add_heading("4.2 WebSocket (Real-time Stream)", level=2)
    doc.add_paragraph("For voice interviews, a persistent WebSocket connection streams audio in both directions.")
    add_code(doc, "Browser <--WebSocket--> Backend (port 8080) <--WebSocket--> OpenAI Realtime API")

    doc.add_heading("5. System Architecture Diagram", level=1)
    add_code(doc,
        "+-------------------------------------------------+\n"
        "|              USER'S BROWSER                      |\n"
        "|  +---------------------------------------------+ |\n"
        "|  |   React Frontend (Vite Dev / Static)        | |\n"
        "|  |   - Pages, Components, Hooks                | |\n"
        "|  |   - AuthContext (JWT in localStorage)        | |\n"
        "|  |   - Axios -> REST API calls                 | |\n"
        "|  |   - WebSocket -> Voice streaming            | |\n"
        "|  |   - MediaPipe -> Proctor gaze tracking      | |\n"
        "|  +------------+--------------+-----------------+ |\n"
        "|               |REST          |WebSocket          |\n"
        "+-------------------------------------------------+\n"
        "                |              |\n"
        "                v              v\n"
        "+-------------------------------------------------+\n"
        "|        SPRING BOOT BACKEND (Java 21)            |\n"
        "|  +----------+  +----------+  +---------+        |\n"
        "|  |Controllers|  | Security |  |WebSocket|        |\n"
        "|  |(REST API) |  |(JWT+CORS)|  | Handler |        |\n"
        "|  +-----+-----+  +----------+  +----+----+       |\n"
        "|        |                            |            |\n"
        "|  +-----v-----+  +----------+  +----v----+       |\n"
        "|  |  Services  |  |CV Parser |  | OpenAI  |       |\n"
        "|  |(Auth,      |  |(PDFBox + |  |Realtime |       |\n"
        "|  |Interview)  |  | GPT-mini)|  |  Proxy  |       |\n"
        "|  +-----+------+  +----------+  +---------+       |\n"
        "|        |                                         |\n"
        "|  +-----v--------------------------------------+  |\n"
        "|  |    JPA + Hibernate (ORM Layer)              |  |\n"
        "|  +-----+--------------------------------------+  |\n"
        "+-------------------------------------------------+\n"
        "         |\n"
        "         v\n"
        "+------------------+    +------------------+\n"
        "|   PostgreSQL DB  |    |   OpenAI APIs    |\n"
        "|  - users         |    |  - Chat (GPT-4o) |\n"
        "|  - user_profiles |    |  - Whisper (STT) |\n"
        "|  - interview_    |    |  - TTS           |\n"
        "|    profiles      |    |  - Realtime      |\n"
        "+------------------+    +------------------+"
    )

    doc.add_heading("6. Project Structure", level=1)
    doc.add_heading("6.1 Backend", level=2)
    add_code(doc,
        "groomguru-service/\n"
        "  src/main/java/com/groomguru/\n"
        "    GroomGuruApplication.java    # Entry point\n"
        "    config/                      # WebSocket + Web config\n"
        "    controller/                  # REST API endpoints\n"
        "    dto/request/                 # Incoming data shapes\n"
        "    dto/response/                # Outgoing data shapes\n"
        "    entity/                      # Database table models\n"
        "    exception/                   # Error handling\n"
        "    handler/                     # WebSocket handler\n"
        "    repository/                  # Database queries\n"
        "    security/                    # JWT + Auth filter\n"
        "    service/impl/                # Business logic\n"
        "    util/                        # Validation helpers\n"
        "  src/main/resources/\n"
        "    application.properties       # Config\n"
        "    db/migration/V1__baseline.sql # DB schema\n"
        "  pom.xml                        # Dependencies"
    )
    doc.add_heading("6.2 Frontend", level=2)
    add_code(doc,
        "groomguru-ui/\n"
        "  src/\n"
        "    main.tsx              # App entry point\n"
        "    App.tsx               # Routing setup\n"
        "    index.css             # Tailwind + custom CSS\n"
        "    components/           # Reusable UI pieces\n"
        "      auth/               # Login, Register, Guards\n"
        "      chat/               # Chat interview screens\n"
        "      voice/              # Voice interview screens\n"
        "      profile/            # Profile page cards\n"
        "      layout/             # Sidebar, Topbar\n"
        "      shared/             # Report, Analysis cards\n"
        "      common/             # CountryCodePicker\n"
        "    pages/                # Page-level components\n"
        "    hooks/voice/          # Audio/voice hooks\n"
        "    services/             # API call functions\n"
        "    context/              # AuthContext (global state)\n"
        "    types/                # TypeScript interfaces\n"
        "    utils/                # Helper functions\n"
        "    constants/            # Config values\n"
        "  package.json            # Dependencies\n"
        "  vite.config.ts          # Build config"
    )

    save(doc, output_dir, "02_Tech_Stack_and_Architecture.docx")


def doc3(output_dir, be_path, fe_path):
    doc = styled_doc("Database Design")

    doc.add_heading("1. What is a Database?", level=1)
    doc.add_paragraph(
        "A database is where all the application's data is permanently stored. When you create an account, "
        "update your profile, or upload a CV, that information is saved in the database. Even if the server "
        "restarts, your data is safe."
    )
    doc.add_paragraph("GroomGuru uses PostgreSQL, a powerful open-source relational database.")

    doc.add_heading("2. Tables Overview", level=1)
    add_table(doc, ["Table Name", "Purpose", "Rows Represent"], [
        ["users", "Authentication data", "One row per registered user"],
        ["user_profiles", "Personal information", "One row per user (name, phone, avatar)"],
        ["interview_profiles", "AI interview context", "One row per user (role, skills, CV)"],
        ["flyway_schema_history", "Migration tracking (system)", "One row per applied migration"],
    ])

    doc.add_heading("3. Table: users", level=1)
    doc.add_paragraph("Stores the core authentication data — everything needed to identify and log in a user.")
    add_table(doc, ["Column", "Type", "Nullable?", "Description"], [
        ["id", "UUID", "No (PK)", "Unique identifier, randomly generated"],
        ["email", "VARCHAR(255)", "No (Unique)", "User's email, used as login identifier"],
        ["password", "VARCHAR(255)", "Yes", "BCrypt-hashed password. NULL for Google users"],
        ["auth_provider", "VARCHAR(50)", "No", "'LOCAL' or 'GOOGLE'"],
        ["provider_id", "VARCHAR(255)", "Yes", "Google's unique user ID"],
        ["deleted", "BOOLEAN", "No (default FALSE)", "Soft delete flag"],
        ["deleted_at", "TIMESTAMP", "Yes", "When soft-deleted"],
        ["last_login_at", "TIMESTAMP", "Yes", "Last successful login time"],
        ["created_at", "TIMESTAMP", "No", "Account creation time"],
        ["updated_at", "TIMESTAMP", "Yes", "Last modification time"],
    ])

    doc.add_heading("4. Table: user_profiles", level=1)
    doc.add_paragraph("Stores personal information, separated from auth data following Single Responsibility Principle.")
    add_table(doc, ["Column", "Type", "Nullable?", "Description"], [
        ["id", "UUID", "No (PK)", "Unique identifier"],
        ["user_id", "UUID", "No (FK, Unique)", "Links to users.id (1:1)"],
        ["first_name", "VARCHAR(255)", "Yes", "Auto-capitalized by backend"],
        ["last_name", "VARCHAR(255)", "Yes", "Auto-capitalized by backend"],
        ["phone_number", "VARCHAR(50)", "Yes", "Stored as '+91 8617015319'"],
        ["avatar_url", "TEXT", "Yes", "URL, gradient name, or base64 data URL"],
        ["created_at", "TIMESTAMP", "No", "Profile creation time"],
        ["updated_at", "TIMESTAMP", "Yes", "Shown as 'Last updated' in UI"],
    ])

    doc.add_heading("5. Table: interview_profiles", level=1)
    doc.add_paragraph("Stores information that helps the AI conduct better, more personalized interviews.")
    add_table(doc, ["Column", "Type", "Nullable?", "Description"], [
        ["id", "UUID", "No (PK)", "Unique identifier"],
        ["user_id", "UUID", "No (FK, Unique)", "Links to users.id (1:1)"],
        ["curr_role", "VARCHAR(255)", "Yes", "Current job title"],
        ["experience_level", "VARCHAR(50)", "Yes", "FRESHER/JUNIOR/MID/SENIOR/LEAD/PRINCIPAL"],
        ["years_of_experience", "VARCHAR(50)", "Yes", "Range: 0-1, 1-3, 3-5, 5-10, 10+"],
        ["target_role", "VARCHAR(255)", "Yes", "Role interviewing for"],
        ["skills", "TEXT", "Yes", "Comma-separated skills"],
        ["industry", "VARCHAR(255)", "Yes", "Technology, Finance, etc."],
        ["bio", "TEXT", "Yes", "Professional summary"],
        ["preferred_interview_type", "VARCHAR(100)", "Yes", "Technical, Behavioral, etc."],
        ["cv_file_name", "VARCHAR(255)", "Yes", "Original CV filename"],
        ["cv_data", "TEXT", "Yes", "Base64-encoded CV PDF"],
        ["cv_uploaded_at", "TIMESTAMP", "Yes", "CV upload time"],
        ["created_at", "TIMESTAMP", "No", "Creation time"],
        ["updated_at", "TIMESTAMP", "Yes", "Last modification time"],
    ])

    doc.add_heading("6. Relationships", level=1)
    add_code(doc,
        "users (1) ---> (1) user_profiles      [ON DELETE CASCADE]\n"
        "users (1) ---> (1) interview_profiles  [ON DELETE CASCADE]"
    )
    doc.add_paragraph("Each user has exactly ONE user_profile and ONE interview_profile.")

    doc.add_heading("7. Indexes", level=1)
    add_table(doc, ["Index Name", "Table", "Column(s)", "Purpose"], [
        ["idx_user_profiles_user_id", "user_profiles", "user_id", "Fast profile lookup by user"],
        ["idx_interview_profiles_user_id", "interview_profiles", "user_id", "Fast interview profile lookup"],
        ["idx_users_email_not_deleted", "users", "email WHERE deleted=FALSE", "Fast active user email lookup"],
    ])

    doc.add_heading("8. Flyway Migration System", level=1)
    doc.add_paragraph(
        "When the database structure needs to change, you create a numbered SQL file (e.g., V2__add_column.sql) "
        "in db/migration/. Flyway applies unapplied migrations in order on startup."
    )

    migration_path = os.path.join(be_path, 'src/main/resources/db/migration')
    if os.path.isdir(migration_path):
        migrations = sorted(os.listdir(migration_path))
        doc.add_paragraph(f"Current migrations ({len(migrations)}):")
        for m in migrations:
            doc.add_paragraph(m, style='List Bullet')

    doc.add_heading("9. Soft Delete", level=1)
    doc.add_paragraph(
        "When a user deletes their account, deleted=TRUE is set instead of removing data. "
        "All queries filter by deleted=FALSE, making soft-deleted users invisible."
    )

    save(doc, output_dir, "03_Database_Design.docx")


def doc4(output_dir, be_path, fe_path):
    doc = styled_doc("API Reference")

    doc.add_heading("1. What is an API?", level=1)
    doc.add_paragraph(
        "An API (Application Programming Interface) is how the frontend and backend communicate. "
        "The frontend sends HTTP requests to endpoints on the backend, and gets JSON responses back."
    )

    doc.add_heading("2. Base URL", level=1)
    doc.add_paragraph("Local: http://localhost:8080")
    doc.add_paragraph("Production: https://groomguru-service-production.up.railway.app")

    doc.add_heading("3. Authentication", level=1)
    doc.add_paragraph("Most endpoints require a JWT token in the header:")
    add_code(doc, 'Authorization: Bearer eyJhbGciOiJIUzM4NCJ9...')

    doc.add_heading("4. Public Endpoints (No Auth)", level=1)

    doc.add_heading("4.1 POST /api/auth/signup", level=2)
    doc.add_paragraph("Create a new account.")
    add_code(doc, '{\n  "firstName": "Ankit",\n  "lastName": "Thapa",\n  "email": "ankit@gmail.com",\n  "password": "SecurePass1!"\n}')

    doc.add_heading("4.2 POST /api/auth/login", level=2)
    add_code(doc, '{\n  "email": "ankit@gmail.com",\n  "password": "SecurePass1!"\n}')

    doc.add_heading("4.3 POST /api/auth/google", level=2)
    add_code(doc, '{\n  "idToken": "eyJhbGciOiJSUzI1NiIs..."\n}')

    doc.add_heading("5. Protected Endpoints (Auth Required)", level=1)

    endpoints = [
        ("GET", "/api/auth/me", "Get current user profile", "None", "UserDto"),
        ("PUT", "/api/auth/profile", "Update profile (name, email, phone)", "UpdateProfileRequest", "UserDto"),
        ("PUT", "/api/auth/avatar", "Update avatar", '{"avatarUrl": "..."}', "UserDto"),
        ("GET", "/api/auth/interview-profile", "Get interview profile", "None", "InterviewProfileResponse"),
        ("PUT", "/api/auth/interview-profile", "Update interview profile", "InterviewProfileRequest", "InterviewProfileResponse"),
        ("PUT", "/api/auth/interview-profile/cv", "Upload CV (base64)", '{"fileName":"...", "cvData":"..."}', "InterviewProfileResponse"),
        ("GET", "/api/auth/interview-profile/cv", "Download CV as PDF", "None", "PDF binary"),
        ("DELETE", "/api/auth/interview-profile/cv", "Delete CV", "None", "InterviewProfileResponse"),
        ("POST", "/api/auth/interview-profile/parse-cv", "AI auto-fill from CV", "None", "InterviewProfileResponse"),
        ("PUT", "/api/auth/password", "Change/set password", "ChangePasswordRequest", "204 No Content"),
        ("DELETE", "/api/auth/account", "Soft-delete account", "None", "204 No Content"),
    ]
    add_table(doc, ["Method", "Path", "Description", "Request", "Response"],
              [[m, p, d, r, resp] for m, p, d, r, resp in endpoints])

    doc.add_heading("6. Interview Endpoints", level=1)
    interview_endpoints = [
        ("POST", "/api/interview/start", "Start interview", "StartInterviewRequest", "StartInterviewResponse"),
        ("POST", "/api/interview/answer", "Submit voice answer", "multipart (audio + params)", "AnswerResponse"),
        ("POST", "/api/interview/chat-answer", "Submit chat answer", "ChatAnswerRequest", "ChatAnswerResponse"),
        ("POST", "/api/interview/upload-cv", "Analyze CV for interview", "multipart (cvFile)", "CvAnalysisResponse"),
        ("POST", "/api/interview/end", "End interview, get report", "EndInterviewRequest", "EndInterviewResponse"),
    ]
    add_table(doc, ["Method", "Path", "Description", "Request", "Response"],
              [[m, p, d, r, resp] for m, p, d, r, resp in interview_endpoints])

    doc.add_heading("7. WebSocket", level=1)
    doc.add_paragraph("ws://localhost:8080/ws/realtime?token=<JWT>")
    doc.add_paragraph("Real-time voice interview via OpenAI Realtime API proxy.")

    doc.add_heading("8. Error Format", level=1)
    add_code(doc, '{"error": "Human-readable error message"}')
    add_table(doc, ["Status", "Meaning"], [
        ["400", "Invalid input"],
        ["401", "Missing/invalid JWT"],
        ["500", "Server error"],
    ])

    save(doc, output_dir, "04_API_Reference.docx")


def doc5(output_dir, be_path, fe_path):
    doc = styled_doc("Frontend Guide")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "The GroomGuru frontend is a Single Page Application (SPA) built with React. The browser loads one HTML page, "
        "and JavaScript dynamically updates the content as you navigate."
    )

    doc.add_heading("2. Entry Point", level=1)
    doc.add_paragraph("When the app loads: index.html -> main.tsx -> wraps in BrowserRouter + GoogleOAuthProvider + AuthProvider -> App.tsx renders routes.")

    doc.add_heading("3. Routing", level=1)
    add_table(doc, ["URL", "Component", "Layout", "Description"], [
        ["/login", "LoginScreen", "None", "Login form"],
        ["/signup", "RegisterScreen", "None", "Registration"],
        ["/", "HomePage", "Sidebar+Topbar", "Dashboard"],
        ["/practice", "PracticePage", "Sidebar+Topbar", "Choose interview mode"],
        ["/setup/voice", "VoiceSetupScreen", "Sidebar+Topbar", "Configure voice interview"],
        ["/setup/chat", "ChatSetupScreen", "Sidebar+Topbar", "Configure chat interview"],
        ["/interview/voice", "RealtimeVoiceInterviewScreen", "Full screen", "Live voice interview"],
        ["/interview/chat", "ChatInterviewScreen", "Full screen", "Chat interview"],
        ["/report", "ReportScreen", "Full screen", "Post-interview report"],
        ["/history", "HistoryPage", "Sidebar+Topbar", "Interview history"],
        ["/profile", "ProfilePage", "Sidebar+Topbar", "Profile management"],
    ])

    doc.add_heading("4. AuthContext — Global State", level=1)
    doc.add_paragraph("Provides: user, token, isLoading, login(), register(), googleLogin(), updateUser(), logout().")
    doc.add_paragraph("Stores JWT in localStorage. Calls GET /api/auth/me on startup to validate saved token.")

    doc.add_heading("5. Services Layer", level=1)
    doc.add_paragraph("http.ts: Axios instance with JWT header and 401 interceptor.")
    doc.add_paragraph("authApi.ts: All /api/auth/* calls.")
    doc.add_paragraph("api.ts: All /api/interview/* calls.")

    doc.add_heading("6. Hooks", level=1)
    for name, desc in [
        ("useCountryCodes", "Fetches country calling codes from restcountries.com, caches in memory"),
        ("useIdleTimer", "15-min idle detection with 60s warning before auto-logout"),
        ("useInterviewTimer", "Countdown timer for interviews"),
        ("useRealtimeVoice", "WebSocket + PCM audio streaming + AI audio playback"),
        ("useAudioRecorder", "Microphone recording via Web Audio API"),
        ("useGazeTracker", "MediaPipe face landmark gaze detection for Proctor Mode"),
        ("useLiveTranscript", "Web Speech API live transcription"),
        ("useSpeechSynthesis", "Base64 MP3 to audio playback"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{name}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("7. Components", level=1)
    component_dir = os.path.join(fe_path, 'src/components')
    if os.path.isdir(component_dir):
        for folder in sorted(os.listdir(component_dir)):
            folder_path = os.path.join(component_dir, folder)
            if os.path.isdir(folder_path):
                files = [f for f in os.listdir(folder_path) if f.endswith('.tsx')]
                doc.add_heading(f"7.x {folder}/", level=2)
                for f in sorted(files):
                    doc.add_paragraph(f, style='List Bullet')

    doc.add_heading("8. Utilities", level=1)
    for name, desc in [
        ("validation.ts", "Password rules, email/name/phone validators, parsePhone"),
        ("speechAnalysis.ts", "WPM, filler words, vocabulary diversity"),
        ("botSignalAnalysis.ts", "Bot-like signal detection in responses"),
        ("date.ts", "Date formatting"),
        ("string.ts", "capitalize() for names"),
        ("image.ts", "Image resizing and base64 conversion"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{name}: ")
        run.font.bold = True
        p.add_run(desc)

    save(doc, output_dir, "05_Frontend_Guide.docx")


def doc6(output_dir, be_path, fe_path):
    doc = styled_doc("Backend Guide")

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("The GroomGuru backend is a Spring Boot 4 application written in Java 21.")

    doc.add_heading("2. Package Structure", level=1)
    add_table(doc, ["Package", "Purpose", "Key Classes"], [
        ["config/", "Configuration", "WebConfig, WebSocketConfig"],
        ["controller/", "REST endpoints", "AuthController, InterviewController"],
        ["dto/request/", "Incoming data", "LoginRequest, RegisterRequest, etc."],
        ["dto/response/", "Outgoing data", "AuthResponse, InterviewProfileResponse, etc."],
        ["entity/", "DB models", "User, UserProfile, InterviewProfile"],
        ["exception/", "Error handling", "GlobalExceptionHandler"],
        ["handler/", "WebSocket", "RealtimeInterviewHandler"],
        ["repository/", "DB access", "UserRepository, etc."],
        ["security/", "Auth", "SecurityConfig, JwtTokenProvider, JwtAuthFilter"],
        ["service/impl/", "Business logic", "AuthServiceImpl, InterviewServiceImpl"],
        ["util/", "Helpers", "ValidationUtils"],
    ])

    doc.add_heading("3. Request Flow", level=1)
    add_code(doc,
        "HTTP Request -> JwtAuthFilter (validate JWT)\n"
        "            -> Controller (route to method)\n"
        "            -> Service (business logic)\n"
        "            -> Repository (JPA query)\n"
        "            -> PostgreSQL\n"
        "            -> Response DTO -> JSON to frontend"
    )

    doc.add_heading("4. Security", level=1)
    doc.add_paragraph("SecurityConfig: CSRF disabled, stateless sessions, CORS, JWT filter.")
    doc.add_paragraph("JwtTokenProvider: HMAC-SHA384, 7-day expiry, contains userId + email.")
    doc.add_paragraph("JwtAuthFilter: Extracts/validates JWT, sets SecurityContext.")

    doc.add_heading("5. AuthServiceImpl (~520 lines)", level=1)
    for name, desc in [
        ("register()", "Creates user + profile, BCrypt password, handles reactivation"),
        ("login()", "Validates credentials, updates lastLoginAt"),
        ("googleAuth()", "Verifies Google token, creates/updates user"),
        ("updateProfile()", "Updates name, email, phone with validation"),
        ("parseCvAndPopulate()", "PDFBox text extraction -> GPT-4o-mini parsing -> auto-fill"),
        ("changePassword()", "Validates strength, BCrypt hashing"),
        ("softDeleteAccount()", "Sets deleted=true"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{name}: ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_heading("6. InterviewServiceImpl", level=1)
    doc.add_paragraph("Uses Spring AI ChatClient for questions/answers, Whisper for STT, TTS for voice.")

    doc.add_heading("7. CvParserService", level=1)
    doc.add_paragraph("Step 1: PDFBox extracts text. Step 2: GPT-4o-mini parses into structured JSON.")

    doc.add_heading("8. RealtimeInterviewHandler", level=1)
    doc.add_paragraph("WebSocket at /ws/realtime — bridges browser audio to OpenAI Realtime API.")

    doc.add_heading("9. InterviewPromptBuilder (400+ lines)", level=1)
    doc.add_paragraph("Builds system prompts adapted to role, experience, interview type, and simulation mode.")

    save(doc, output_dir, "06_Backend_Guide.docx")


def doc7(output_dir, be_path, fe_path):
    doc = styled_doc("Local Setup & Run Guide")

    doc.add_heading("1. Prerequisites", level=1)
    add_table(doc, ["Software", "Version", "Check Command"], [
        ["Java JDK", "21+", "java -version"],
        ["Node.js", "18+ (22 recommended)", "node -v"],
        ["PostgreSQL", "14+", "psql --version"],
        ["Git", "Any recent", "git --version"],
    ])

    doc.add_heading("2. Clone Repositories", level=1)
    add_code(doc, "git clone https://github.com/Ankit-mangar/groomguru-service.git\ngit clone https://github.com/Ankit-mangar/groomguru-ui.git")

    doc.add_heading("3. Set Up PostgreSQL", level=1)
    add_code(doc, "psql\nCREATE DATABASE groomguru;\n\\q")
    doc.add_paragraph("Flyway creates all tables automatically on first backend start.")

    doc.add_heading("4. Backend Environment (.env)", level=1)
    add_code(doc, "# groomguru-service/.env\nOPENAI_API_KEY=sk-your-key\nJWT_SECRET=random-64-char-string\nGOOGLE_CLIENT_ID=your-client-id.apps.googleusercontent.com")
    add_note(doc, "Never commit .env to Git!")

    doc.add_heading("5. Frontend Environment (.env)", level=1)
    add_code(doc, "# groomguru-ui/.env\nVITE_API_BASE_URL=http://localhost:8080\nVITE_WS_BASE_URL=ws://localhost:8080/ws/realtime\nVITE_GOOGLE_CLIENT_ID=your-client-id.apps.googleusercontent.com")

    doc.add_heading("6. Run Backend", level=1)
    add_code(doc, "cd groomguru-service\nexport $(grep -v '^#' .env | xargs) && ./mvnw spring-boot:run")

    doc.add_heading("7. Run Frontend", level=1)
    add_code(doc, "cd groomguru-ui\nnpm install  # first time only\nnpm run dev")

    doc.add_heading("8. Verify", level=1)
    for i, step in enumerate([
        "Open http://localhost:5173/ — see login page",
        "Register with email and password",
        "Go to Profile — update name",
        "Practice Interview — start a chat interview",
    ], 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')

    doc.add_heading("9. Troubleshooting", level=1)
    for problem, fix in [
        ("JWT_SECRET not found", "Run: export $(grep -v '^#' .env | xargs) before mvnw"),
        ("DB connection refused", "Ensure PostgreSQL is running"),
        ("CORS errors", "Check ports: BE=8080, FE=5173"),
        ("Stale frontend", "Kill Vite, rm -rf node_modules/.vite, restart, hard refresh"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{problem}: ")
        run.font.bold = True
        p.add_run(fix)

    save(doc, output_dir, "07_Local_Setup_and_Run_Guide.docx")


def doc8(output_dir, be_path, fe_path):
    doc = styled_doc("Deployment Guide")

    doc.add_heading("1. What is Deployment?", level=1)
    doc.add_paragraph("Putting your app on a cloud server with a public URL. GroomGuru uses Railway (railway.app).")

    doc.add_heading("2. Railway Architecture", level=1)
    add_table(doc, ["Service", "Type", "URL"], [
        ["groomguru-service", "Java Web Service", "groomguru-service-production.up.railway.app"],
        ["groomguru-ui", "Static Site", "groomguru-ui-production.up.railway.app"],
        ["PostgreSQL", "Database Plugin", "Internal connection"],
    ])

    doc.add_heading("3. Backend Setup", level=1)
    doc.add_paragraph("1. New Project -> Deploy from GitHub -> groomguru-service")
    doc.add_paragraph("2. Add PostgreSQL plugin")
    doc.add_paragraph("3. Set environment variables:")
    add_table(doc, ["Variable", "Source"], [
        ["PGHOST/PGPORT/PGDATABASE/PGUSER/PGPASSWORD", "From PostgreSQL Connect tab"],
        ["OPENAI_API_KEY", "Your OpenAI key"],
        ["JWT_SECRET", "Your JWT secret"],
        ["GOOGLE_CLIENT_ID", "Your Google OAuth ID"],
        ["CORS_ALLOWED_ORIGINS", "Frontend Railway URL"],
        ["PORT", "8080"],
    ])

    doc.add_heading("4. Frontend Setup", level=1)
    doc.add_paragraph("1. Add GitHub repo -> groomguru-ui")
    doc.add_paragraph("2. Set environment variables:")
    add_table(doc, ["Variable", "Value"], [
        ["VITE_API_BASE_URL", "https://groomguru-service-production.up.railway.app"],
        ["VITE_WS_BASE_URL", "wss://groomguru-service-production.up.railway.app/ws/realtime"],
        ["VITE_GOOGLE_CLIENT_ID", "Your Google OAuth ID"],
    ])

    doc.add_heading("5. Flyway in Production", level=1)
    doc.add_paragraph("Runs automatically on deploy. Never modify existing migration files. Always create new ones.")

    doc.add_heading("6. Continuous Deployment", level=1)
    doc.add_paragraph("Push to main -> Railway auto-builds -> zero-downtime deploy -> health check.")

    doc.add_heading("7. Google OAuth for Production", level=1)
    doc.add_paragraph("Add production frontend URL to Authorized JavaScript origins in Google Cloud Console.")

    doc.add_heading("8. Troubleshooting", level=1)
    for problem, fix in [
        ("Build fails", "Check build logs for missing env vars"),
        ("CORS errors", "Verify CORS_ALLOWED_ORIGINS matches frontend URL exactly"),
        ("Schema errors", "Create new Flyway migration, never edit existing"),
        ("WebSocket fails", "Use wss:// not ws:// in production"),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{problem}: ")
        run.font.bold = True
        p.add_run(fix)

    save(doc, output_dir, "08_Deployment_Guide.docx")


# ── Main ─────────────────────────────────────────────

if __name__ == "__main__":
    args = parse_args()
    os.makedirs(args.output, exist_ok=True)

    be = os.path.abspath(args.be_path)
    fe = os.path.abspath(args.fe_path)
    out = os.path.abspath(args.output)

    print(f"Backend path:  {be}")
    print(f"Frontend path: {fe}")
    print(f"Output dir:    {out}")
    print()

    doc1(out, be, fe)
    doc2(out, be, fe)
    doc3(out, be, fe)
    doc4(out, be, fe)
    doc5(out, be, fe)
    doc6(out, be, fe)
    doc7(out, be, fe)
    doc8(out, be, fe)

    print(f"\nDone! All documents saved to: {out}/")
